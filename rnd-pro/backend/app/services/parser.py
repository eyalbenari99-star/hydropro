"""Real spec text extractor: PDF, Word, Excel, plain text, markdown, CSV, JSON.

Optional deps with graceful fallback:
  - pypdf for PDF
  - python-docx for Word
  - openpyxl for Excel (already required for BOM)
Returns (text, meta) where meta = {maybeCabinet, maybePlumbing, format, sourceExt, lengthChars}.
"""
import io, re
from typing import Dict, Any, Tuple


CABINET_KEYWORDS = ['MCB', 'MCCB', 'NXB', 'NXC', 'NM1', 'JR36', 'breaker', 'contactor',
                    'cabinet', 'feeder', 'busbar', 'kW pump', 'kW motor', 'pole',
                    'curve C', 'curve D', 'main breaker', 'overload', 'IP54', 'IP65']
PLUMB_KEYWORDS  = ['m3/h', 'm^3/h', 'pump', 'filter', 'zone', 'pressure', 'bar',
                   'fertigation', 'irrigation', 'dripper', 'DN', 'PVC', 'PE pipe',
                   'flow rate', 'manifold', 'dosing', 'EC sensor', 'pH sensor']


def extract_text_and_meta(filename: str, content: bytes) -> Tuple[str, Dict[str, Any]]:
    ext = (filename.rsplit('.', 1)[-1] if '.' in filename else '').lower()
    text, fmt = '', 'unknown'

    if ext in ('txt', 'md', 'csv', 'tsv'):
        text = content.decode('utf-8', errors='replace')
        fmt = 'flat'
    elif ext == 'json':
        text = content.decode('utf-8', errors='replace')
        fmt = 'json'
    elif ext == 'pdf':
        text, fmt = _pdf(content)
    elif ext in ('docx', 'doc'):
        text, fmt = _docx(content)
    elif ext in ('xlsx', 'xlsm', 'xls'):
        text, fmt = _xlsx(content)
    elif ext in ('png', 'jpg', 'jpeg', 'gif', 'webp'):
        text, fmt = '', 'image-needs-ocr'

    if '|' in text and re.search(r'\|.*\|.*\|', text) and re.search(r'\|[\s\-:]+\|', text):
        fmt = 'markdown-table'

    low = text.lower()
    maybe_cab  = sum(1 for k in CABINET_KEYWORDS if k.lower() in low) >= 2
    maybe_plumb = sum(1 for k in PLUMB_KEYWORDS if k.lower() in low) >= 2

    return text, {
        'maybeCabinet': maybe_cab,
        'maybePlumbing': maybe_plumb,
        'format': fmt,
        'sourceExt': ext,
        'lengthChars': len(text),
    }


def _pdf(content: bytes) -> Tuple[str, str]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        parts = []
        for page in reader.pages:
            try: parts.append(page.extract_text() or '')
            except Exception: pass
        return '\n\n'.join(parts), 'pdf-text'
    except ImportError:
        return '[PDF parsing requires pypdf - pip install pypdf]', 'pdf-error'
    except Exception as e:
        return f'[PDF parse error: {e}]', 'pdf-error'


def _docx(content: bytes) -> Tuple[str, str]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(' | '.join(c.text for c in row.cells))
        return '\n'.join(parts), 'word-text'
    except ImportError:
        return '[Word parsing requires python-docx - pip install python-docx]', 'word-error'
    except Exception as e:
        return f'[Word parse error: {e}]', 'word-error'


def _xlsx(content: bytes) -> Tuple[str, str]:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f'## Sheet: {sheet}')
            for row in ws.iter_rows(values_only=True):
                parts.append(' | '.join('' if v is None else str(v) for v in row))
            parts.append('')
        return '\n'.join(parts), 'xlsx-table'
    except ImportError:
        return '[Excel parsing requires openpyxl]', 'xlsx-error'
    except Exception as e:
        return f'[Excel parse error: {e}]', 'xlsx-error'


# ---------- Deterministic spec parsers ----------

def parse_cabinet_text(text: str, format_hint: str = 'unknown'):
    """Parse plain text into a minimal CabinetSpecV2 by scanning for kW patterns."""
    from app.models.cabinet import CabinetSpecV2, Section, Rail, DeviceV2, CabinetMeta
    items = []
    for i, line in enumerate(text.splitlines()):
        m = re.search(r'(\d+(?:\.\d+)?)\s*kW\s+(\w+)', line, re.IGNORECASE)
        if m:
            kw = float(m.group(1))
            kind = m.group(2).lower()
            ref_idx = len(items) + 1
            items.append(DeviceV2(
                id=f'F{ref_idx}', sym='MCB 3P', kind='breaker',
                label=f'{kw}kW {kind}', ref=f'{ref_idx}QF',
                poles=3, rating=f'{int(kw*3)} A',
                manufacturer='CHINT', partNo=f'NXB-63 D{int(kw*3)} 3P'
            ))
    sections = []
    if items:
        sections.append(Section(id='SEC_FEEDERS', label='Imported feeders',
                                type='power', color='#2563eb',
                                rails=[Rail(id='R_FB', label='Feeder MCBs', railType='DIN', items=items)]))
    return CabinetSpecV2(title='Imported cabinet', subtitle=f'{len(items)} feeders parsed',
                         sections=sections, cabinet=CabinetMeta())


def parse_plumbing_text(text: str, format_hint: str = 'unknown'):
    """Parse plain text into a minimal PlumbingSpec by scanning for flow/pressure."""
    from app.models.plumbing import PlumbingSpec, PlumbNode, PlumbEdge
    fm = re.search(r'(\d+(?:\.\d+)?)\s*m3?/h', text, re.IGNORECASE)
    pm = re.search(r'(\d+(?:\.\d+)?)\s*bar', text, re.IGNORECASE)
    zm = re.search(r'(\d+)\s*zone', text, re.IGNORECASE)
    flow = float(fm.group(1)) if fm else 10.0
    p    = float(pm.group(1)) if pm else 3.0
    zones = int(zm.group(1)) if zm else 4
    nodes = [
        PlumbNode(id='SRC',  type='source', label='Source',           flowM3h=flow, pressureBar=p),
        PlumbNode(id='PUMP', type='pump',   label=f'Pump {flow}m3/h x {p}bar'),
        PlumbNode(id='FLT',  type='filter', label='Filter 130um'),
        PlumbNode(id='MAN',  type='manifold', label=f'Manifold {zones} ports'),
    ]
    edges = [
        PlumbEdge(fromId='SRC',  toId='PUMP'),
        PlumbEdge(fromId='PUMP', toId='FLT'),
        PlumbEdge(fromId='FLT',  toId='MAN'),
    ]
    for z in range(zones):
        zid = f'Z{z+1}'
        nodes.append(PlumbNode(id=f'OUT_{zid}', type='outlet', label=f'Zone {z+1}'))
        edges.append(PlumbEdge(fromId='MAN', toId=f'OUT_{zid}'))
    return PlumbingSpec(title='Imported hydraulic system',
                        zones=zones, nodes=nodes, edges=edges)
